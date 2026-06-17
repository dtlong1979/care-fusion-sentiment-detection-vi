# -*- coding: utf-8 -*-
"""Build a single .docx compiling the CARE-Fusion working draft + findings log."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"C:\Dev\care-fusion-sentiment-detection-vi")
OUT = ROOT / "docs" / "CARE-Fusion_Working_Draft.docx"

doc = Document()
nstyle = doc.styles["Normal"]
nstyle.font.name = "Arial"
nstyle.font.size = Pt(11)
for hs, sz in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)]:
    try:
        st = doc.styles[hs]; st.font.name = "Arial"
        st.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        if hs != "Title":
            st.font.size = Pt(sz)
    except KeyError:
        pass


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)
def b(t): doc.add_paragraph(t, style="List Bullet")
def n(t): doc.add_paragraph(t, style="List Number")


def p(*parts):
    """parts: str (plain) or (str, 'b') tuple for bold; or 'i' for italic."""
    par = doc.add_paragraph()
    for part in parts:
        if isinstance(part, tuple):
            run = par.add_run(part[0])
            if "b" in part[1]:
                run.bold = True
            if "i" in part[1]:
                run.italic = True
        else:
            par.add_run(part)
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htext); r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return t


# ============================== TITLE ==============================
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("CARE-Fusion: Hợp nhất văn bản và biểu tượng cảm xúc theo chế độ tương tác "
               "cho phân loại cảm xúc tiếng Việt")
r.bold = True; r.font.size = Pt(18)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Bản nháp đang thực hiện & Nhật ký phát hiện — cập nhật 17/06/2026")
sr.italic = True; sr.font.size = Pt(11)
repo = doc.add_paragraph(); repo.alignment = WD_ALIGN_PARAGRAPH.CENTER
repo.add_run("Mã nguồn: github.com/dtlong1979/care-fusion-sentiment-detection-vi").font.size = Pt(10)
doc.add_paragraph()
note = doc.add_paragraph()
nr = note.add_run("Ghi chú: Đây là bản nháp nội bộ. Các trích dẫn đánh dấu [cần xác minh] là công trình "
                  "rất mới (2024–2026), cần mở nguồn gốc kiểm tra tác giả/năm/venue trước khi nộp. "
                  "Số liệu CARE-Fusion là 'config cũ' (chưa có bản vá aux/confidence) trừ khi ghi rõ; "
                  "kết quả 5-seed của config mới đang chạy.")
nr.italic = True; nr.font.size = Pt(9.5)
doc.add_paragraph()

# ============================== ABSTRACT ==============================
h1("Tóm tắt")
p("Chúng tôi nghiên cứu vai trò của biểu tượng cảm xúc (emoji/emoticon) trong phân loại cảm xúc "
  "văn bản tiếng Việt trên mạng xã hội. Từ ViGoEmotions, chúng tôi xây dựng một tập dẫn xuất gồm "
  "6.813 bình luận có đồng thời văn bản và biểu tượng cảm xúc, gộp 27 nhãn chi tiết thành 6 nhóm. "
  "Chúng tôi đề xuất CARE-Fusion, một khung hợp nhất văn bản–marker với định tuyến theo ba chế độ "
  "tương tác (dư thừa/bổ sung/xung đột) dựa trên phân bố cảm xúc thực nghiệm của marker (q_j) ước "
  "lượng riêng cho tiếng Việt. Qua thực nghiệm 5-seed có kiểm định thống kê, chúng tôi thấy: "
  "(i) biểu tượng cảm xúc cải thiện đáng kể so với chỉ-văn-bản (+~5 macro-F1); (ii) trên bộ dữ "
  "liệu này, hợp nhất đơn giản đã đạt hiệu năng tương đương các kiến trúc phức tạp hơn; và "
  "(iii) khó khăn cốt lõi nằm ở sự nhập nhằng amusement–sarcasm và tính đa nghĩa của emoji, mà "
  "chúng tôi phân tích định lượng (entropy của q_j, lát mẫu xung đột cực tính). Kiểm định 5-seed "
  "(cả config cũ lẫn config có aux+confidence-gating) cho thấy định tuyến theo chế độ KHÔNG vượt "
  "các baseline hợp nhất đơn giản một cách có ý nghĩa thống kê trên bộ này; do đó chúng tôi định "
  "vị đóng góp theo hướng phân tích/tài nguyên + tính giải thích được. Một phát hiện trọng tâm: "
  "biểu tượng cảm xúc là tín hiệu HAI LƯỠI — giúp ở văn bản đồng thuận nhưng gài bẫy ở mỉa mai, nơi "
  "mô hình chỉ-văn-bản lại kháng nhiễu tốt nhất.")

# ============================== 1. INTRO ==============================
h1("1. Giới thiệu")
p("Trên mạng xã hội tiếng Việt, emoji và emoticon (vd “:))”, “:v”, “🤣”) là một kênh biểu đạt cảm "
  "xúc song song với ngôn ngữ, đôi khi củng cố, đôi khi đảo ngược nghĩa của văn bản (mỉa mai). "
  "Phần lớn nghiên cứu cảm xúc tiếng Việt hiện nay là chỉ-văn-bản; số ít có dùng emoji thì xử lý ở "
  "mức tiền xử lý hoặc trang trí, chưa mô hình hóa vai trò tương tác phân biệt của từng loại "
  "biểu tượng. Bài báo này: (1) xây một tập dữ liệu chuyên biệt emoji+text tiếng Việt; (2) đề xuất "
  "khung CARE-Fusion với định tuyến theo chế độ; (3) phân tích sâu khi nào và vì sao biểu tượng "
  "cảm xúc giúp ích, kèm các phát hiện về tính nhập nhằng của emoji.")

# ============================== 2. RELATED WORKS ==============================
h1("2. Công trình liên quan (Related Works)")
h2("2.1 Emoji như tín hiệu cảm xúc và hợp nhất emoji–văn bản")
p("Novak và cộng sự (2015) xây Emoji Sentiment Ranking — từ điển phân cực cảm xúc đầu tiên cho 751 "
  "emoji. Eisner và cộng sự (2016) đề xuất emoji2vec; Felbo và cộng sự (2017, DeepMoji) dùng emoji "
  "làm giám sát từ xa để học biểu diễn cảm xúc. Wang và cộng sự (2024, EMFSA) hợp nhất emoji–văn bản "
  "bằng cross-attention nhưng chỉ mô hình hóa quan hệ bổ sung và đánh giá trên tiếng Anh. Các hướng "
  "này khai thác emoji ở mức biểu diễn/nhãn, hoặc một đường fusion bổ sung duy nhất.")
h2("2.2 Xung đột emoji–văn bản và phát hiện mỉa mai")
p("Chauhan và cộng sự (2022) đề xuất Gated Emoji-aware Multimodal Attention cho mỉa mai đa thể thức; "
  "Zhu và cộng sự (2025, ECCA) [cần xác minh] dùng contrastive attention để làm nổi bật sự bất nhất "
  "cảm xúc. Các công trình này xử lý xung đột như một đường suy luận riêng cho tác vụ mỉa mai, "
  "không phải một trong nhiều chế độ được định tuyến trong phân loại cảm xúc tổng quát.")
h2("2.3 Cảm xúc tiếng Việt")
p("PhoBERT (Nguyen & Nguyen, 2020) và ViSoBERT (Nguyen và cộng sự, 2023) là các mô hình tiền-huấn "
  "luyện chủ đạo. Các nghiên cứu sentiment/emotion tiếng Việt (khảo sát Thin và cộng sự, 2023; "
  "ensemble 2023) phần lớn chỉ-văn-bản. Công trình gần nhất về emoji trong cảm xúc tiếng Việt "
  "(Tran Thi Dung và cộng sự, 2024) [cần xác minh] chỉ gán emoji như minh họa đầu ra. Bộ dữ liệu "
  "ViGoEmotions (Hung và cộng sự, 2026) [cần xác minh] khảo sát tiền xử lý emoji nhưng coi emoji là "
  "một phần của luồng văn bản.")
h2("2.4 Khoảng trống nghiên cứu")
p("Chưa có mô hình nào hợp nhất emoji+text cho cảm xúc tiếng Việt với định tuyến tường minh theo "
  "ba chế độ tương tác (redundancy/complementarity/conflict). CARE-Fusion nhắm vào khoảng trống này.")

# ============================== 3. DATASET ==============================
h1("3. Xây dựng tập dữ liệu")
p("Tập dữ liệu dẫn xuất từ ViGoEmotions (Hung và cộng sự, 2026) [cần xác minh] — bộ cảm xúc tiếng "
  "Việt với 27 nhãn chi tiết + trung lập, ba phân hoạch train/val/test. Việc tái sử dụng nhãn đã "
  "kiểm định giúp tránh các vấn đề đạo đức, bảo vệ thông tin cá nhân và đảm bảo chất lượng gán nhãn.")
p("Chúng tôi giữ các mẫu thỏa mãn đồng thời: (i) text có nội dung ngôn ngữ tự nhiên; (ii) text chứa "
  "ít nhất một biểu tượng cảm xúc (emoji Unicode hoặc emoticon như “:))”, “:(”, “:v”, “<3”, “^^”, "
  "“T_T”). Tiêu chí này giới hạn có chủ đích vào văn bản có biểu tượng cảm xúc.")
p("Ba phân hoạch gốc được gộp thành một tệp duy nhất nhưng ", ("bảo toàn nguyên vẹn phân hoạch "
  "train/val/test gốc", "b"), " qua trường source_split; mọi thí nghiệm tuân thủ phân hoạch này "
  "(không chia lại) nên không rò rỉ. Sau lọc: 6.813 mẫu (5.442 train, 679 val, 692 test); "
  "4.929 mẫu chỉ chứa emoji Unicode, 1.813 mẫu chỉ chứa emoticon, 71 mẫu chứa cả hai.")
p("27 nhãn được ánh xạ vào 6 nhóm: positive, interest, sadness, anger, fear, neutral (interest là "
  "cụm cảm xúc nhận thức/hóa trị mơ hồ). Với mẫu đa nhãn, áp dụng quy tắc ưu tiên "
  "anger > sadness > fear > positive > interest > neutral. ", ("31,69% mẫu (2.159/6.813) là đa nhóm",
  "b"), " và chịu tác động của quy tắc này (28,56% hai nhóm, 2,99% ba nhóm, 0,13% bốn nhóm); độ nhạy "
  "với quy tắc được kiểm chứng ở Mục 6 (F5). Phân phối nhãn mất cân bằng nặng (Bảng 1) nên dùng "
  "macro-F1 làm độ đo chính.")
p("Bảng 1. Thống kê tập dữ liệu dẫn xuất.", ("", "i"))
table(["", "positive", "sadness", "anger", "interest", "fear", "neutral", "Tổng"],
      [["train", 2623, 1404, 1055, 168, 153, 39, 5442],
       ["val", 289, 206, 132, 24, 24, 4, 679],
       ["test", 345, 179, 124, 17, 21, 6, 692],
       ["Tổng", 3257, 1789, 1311, 209, 198, 49, 6813]])

# ============================== 4. METHOD ==============================
h1("4. Phương pháp")
p("Cho một bình luận gồm văn bản và chuỗi marker cảm xúc, CARE-Fusion dự đoán một trong 6 nhóm. "
  "Mọi tài nguyên dẫn xuất (q_j, nhãn chế độ yếu, đồ thị PMI) chỉ dựng từ tập train.")
h2("4.1 Tài nguyên từ train")
b("q_j: phân bố cảm xúc thực nghiệm của mỗi marker (Laplace smoothing; back-off toàn cục cho marker "
  "tần suất thấp). Mỗi marker là MỘT PHÂN PHỐI trên 6 nhóm — không gán cứng một nhãn.")
b("Đồ thị PPMI giữa marker và nhóm cảm xúc (đồng xuất hiện trên train) cho bước làm giàu bằng GCN.")
b("Nhãn chế độ yếu (B2): p_text out-of-fold (5-fold) so với q_j qua JSD → redundancy/complementarity/"
  "conflict, dùng giám sát yếu bộ định tuyến.")
h2("4.2 Kiến trúc")
b("Bộ mã hóa văn bản: PhoBERT + attention pooling → biểu diễn câu và p_text.")
b("Bộ mã hóa marker: e_j từ q_j (hoặc embedding học nếu tần suất thấp) + cường độ c_j = log(1+đếm).")
b("Chú ý chéo có cổng giữa marker và chuỗi token → g_j.")
b("Bộ định tuyến: r_j = softmax(W[g_j, e_j, δ_j, c_j]) ∈ Δ², với δ_j = JSD(p_text ‖ q_j).")
b("Hợp nhất theo chế độ (cải tiến): mỗi chế độ k trộn dòng marker và dòng text qua cổng β_k "
  "(khởi tạo thiên lệch: redundancy text-nặng, conflict marker-nặng) + tinh chỉnh phi tuyến.")
b("Điều tiết theo độ nhập nhằng (cải tiến): trọng số marker hiệu dụng = β_k · conf_j, với "
  "conf_j = 1 − H(q_j)/log C; emoji nhập nhằng (entropy cao) → tin văn bản hơn.")
b("Làm giàu bằng GCN trên đồ thị PMI → bộ phân loại tuyến tính → ŷ.")
h2("4.3 Hàm mục tiêu")
p("L = L_cls + λ1·L_route + λ2·L_cf + λ_aux·L_aux. Trong đó L_cls là class-balanced focal loss "
  "(Cui và cộng sự, 2019); L_route là CE giám sát yếu cho router; L_cf là ràng buộc nhất quán "
  "phản thực (xóa/đảo marker). ", ("L_aux (cải tiến):", "b"), " giám sát đầu p_text bằng nhãn vàng "
  "để δ_j = JSD(p_text‖q_j) có ý nghĩa và router được giám sát nhất quán. Tối ưu AdamW, "
  "discriminative LR, dừng sớm theo macro-F1 val, ≥5 seed.")

# ============================== 5. SETUP ==============================
h1("5. Thiết lập thí nghiệm")
p("PhoBERT-base; max_length 256 (thực nghiệm chính); AdamW (PhoBERT 2e-5, head 1e-3); 5 seed "
  "{13,42,123,2024,7}; báo cáo mean±std. Baseline tái cài đặt trên cùng dữ liệu/split/seed: "
  "B0 majority, B1 text-only, B2 marker-concat, B3 scalar-gated, B4 cross-attention (no routing). "
  "Ablation của CARE-Fusion: −routing, −δ, −L_cf, −cường độ, −GCN, −confidence, −aux. "
  "Đánh giá: macro-F1 (6 lớp & 5 lớp đông), F1 từng lớp, ma trận nhầm lẫn; phân tầng theo chế độ "
  "(F2); causal probing (F3); faithfulness (F4); độ nhạy quy tắc gộp nhãn (F5); bootstrap CI, "
  "Wilcoxon, McNemar (G).")

# ============================== 6. RESULTS ==============================
h1("6. Kết quả & Phân tích")
h2("6.1 Kết quả chính (A100, 5 seed, config cũ)")
p("Bảng 2. Macro-F1 trên test (5 seed, mean±std). Lát sarcasm = mẫu có cực marker ngược cực nhãn "
  "(145 mẫu).", ("", "i"))
table(["Mô hình", "Overall", "Lát sarcasm"],
      [["B0 majority", "0.111", "—"],
       ["B1 text-only", "0.449 ± 0.034", "0.529 ± 0.058"],
       ["B2 marker-concat", "0.473 ± 0.006", "0.571 ± 0.038"],
       ["B3 scalar-gated", "0.504 ± 0.015", "0.509 ± 0.051"],
       ["B4 cross-attn", "0.496 ± 0.020", "0.446 ± 0.070"],
       ["CARE-Fusion (full)", "0.506 ± 0.033", "0.493 ± 0.055"],
       ["  − routing", "0.506 ± 0.020", "0.482 ± 0.027"],
       ["  − delta", "0.510 ± 0.021", "0.528 ± 0.057"],
       ["  − counterfactual", "0.504 ± 0.020", "0.457 ± 0.025"],
       ["  − intensity", "0.494 ± 0.018", "0.471 ± 0.030"],
       ["  − GCN", "0.517 ± 0.023", "0.512 ± 0.047"]])
p("Nhận xét: (i) marker giúp rõ — mọi mô hình fusion (~0.50) vượt B1 text-only (0.449), +~5 điểm; "
  "(ii) ", ("ở config cũ, CARE-Fusion KHÔNG vượt baseline nào có ý nghĩa thống kê", "b"),
  " — Wilcoxon CARE vs mọi mô hình cho p ≥ 0,31 (overall) và p ≥ 0,12 (sarcasm); các biến thể fusion "
  "hòa nhau ~0.50–0.52; (iii) trên lát sarcasm, B2 concat (0.571) lại cao nhất, gợi ý hợp nhất đơn "
  "giản đã đủ. Đây là kết quả config cũ; config mới (aux+confidence) đang được kiểm 5-seed.")
h2("6.2 Đóng góp của biểu tượng cảm xúc")
p("Chênh lệch B1 (0.449) → các mô hình fusion (~0.50) xác nhận biểu tượng cảm xúc mang thông tin "
  "bổ sung cho cảm xúc tiếng Việt. Causal Sensitivity Score ≈ 0.73 và độ dịch dự đoán KL ≈ 0.45 khi "
  "xóa marker (F3/F4) cho thấy mô hình thực sự dùng marker.")
h2("6.3 Phân tích dữ liệu (đặc trưng tập)")
b("Conflict/mỉa mai: 21,9% mẫu có-marker có cực marker NGƯỢC cực nhãn vàng (train 21,8% / val 24,1% "
  "/ test 21,0%) — đủ tồn tại nhưng là thiểu số.")
b("Nhiễu nhãn do quy tắc gộp: chỉ 2,3% mẫu bị 'lật cực' (POSITIVE-đa số → gán NEGATIVE); 0% chiều "
  "ngược lại → mapping lành, không phải nguồn nhiễu lớn.")
b("Bất đối xứng: trong các ca text–emoji ngược cực, cực NEGATIVE thắng yếu (~53–54%), một phần do "
  "quy tắc ưu tiên âm-trước → tín hiệu giải-quyết-conflict yếu/nhiễu.")
b("Hài hước trên chủ đề tiêu cực (dark humor): nhiều ca 'text âm + emoji cười → nhãn POSITIVE' là "
  "amusement thật (nhãn theo emoji), trong khi 'text âm + :)) → anger' là mỉa mai (nhãn theo text). "
  "Cùng bề mặt, hai cách giải quyết khác nhau → bài toán bản chất khó.")
h2("6.4 Phân tích emoji: tính nhập nhằng (entropy của q_j)")
p("Vì q_j là phân phối học từ dữ liệu, tính nhập nhằng được đo trực tiếp bằng entropy (max=1,79). "
  "Bảng 3 cho thấy một số emoji rất nhập nhằng, và đáng chú ý ", ("“:)” và “🙂” nghiêng GIẬN", "b"),
  " — kiểu dùng mỉa mai đặc thù tiếng Việt mà từ điển phương Tây (😊=vui) sẽ gán sai. Đây là luận "
  "điểm định lượng cho việc dùng q_j thực nghiệm tiếng Việt.")
p("Bảng 3. Phân bố cảm xúc thực nghiệm và entropy của một số marker (freq cao).", ("", "i"))
table(["Marker", "freq", "entropy", "top nhóm"],
      [[":))", 651, "1.08", "positive .58 / anger .29"],
       ["😂", 449, "0.82", "positive .75 / anger .14"],
       [":)", 119, "1.36", "anger .39 / positive .37  (nhập nhằng)"],
       ["🙂", 135, "1.33", "anger .45 / positive .28  (nhập nhằng)"],
       ["😌", 168, "1.36", "positive .36 / sadness .35  (nhập nhằng)"],
       [":v", 100, "1.34", "positive .45 / anger .31  (nhập nhằng)"],
       ["❤️", 134, "0.52", "positive .88  (rõ ràng)"],
       ["😢", 132, "0.69", "sadness .83  (rõ ràng)"],
       ["😔", 113, "0.62", "sadness .86  (rõ ràng)"]])
h2("6.5 Chẩn đoán cơ chế")
b("Neo p_text (L_aux): không có L_aux, đầu p_text không được giám sát → router định tuyến SAI theo "
  "congruence (redundancy ↔ δ cao, ngược ngữ nghĩa). Thêm L_aux → router căn đúng: redundancy δ≈0.05 "
  "< complementarity ≈0.14 < conflict ≈0.21. Đây là cải tiến cơ chế tái lập được.")
b("β theo chế độ (negative result trung thực): khởi tạo thiên lệch [0.27,0.5,0.73] → học ra "
  "[0.281,0.494,0.694] (gần như không đổi); khởi tạo TRUNG TÍNH [0.5,0.5,0.5] → học ra "
  "[0.505,0.491,0.499] (PHẲNG). → Dữ liệu KHÔNG tự đòi hỏi emoji-áp-đảo-ở-conflict; lợi ích (nếu "
  "có) đến từ prior ta áp, không phải đặc tính dữ liệu.")
b("Điều tiết theo entropy (confidence-gating): 1-seed ban đầu rất hứa hẹn (lát sarcasm 0.582), "
  "NHƯNG kiểm 5-seed cho thấy cú nhảy đó KHÔNG giữ (xem Bảng 4) — phần lớn là dao động 1-seed.")
p("Bảng 4. Config MỚI (aux + confidence-gating), 5 seed, mean±std.", ("", "i"))
table(["Mô hình", "Overall", "Lát sarcasm"],
      [["B4 cross-attn (có aux)", "0.5133 ± 0.012", "0.5158 ± 0.052"],
       ["CARE_full (aux + confidence)", "0.5169 ± 0.018", "0.5227 ± 0.072"],
       ["  − confidence", "0.5102 ± 0.024", "0.4462 ± 0.034"]])
p("Wilcoxon: CARE_full vs B4 → overall p=0.44, lát sarcasm p=1.00 (không ý nghĩa). CARE_full vs "
  "−confidence → lát sarcasm p=0.19 (confidence giúp về xu hướng nhưng chưa ý nghĩa; lát 145 mẫu "
  "có std rất lớn). Mốc B2 concat (0.571) vẫn cao hơn CARE_full(conf) trên lát sarcasm.")
h2("6.6 Định vị đóng góp (đã chốt: hướng phân tích trung thực)")
p("Kiểm định 5-seed bác bỏ claim “định tuyến theo chế độ vượt hợp nhất đơn giản” trên bộ này. "
  "Do đó bài định vị quanh:")
b("Biểu tượng cảm xúc cải thiện đáng kể nhận diện cảm xúc tiếng Việt (fusion ~0.50 vs text-only 0.45).")
b("Phát hiện hữu ích/cảnh báo: hợp nhất ĐƠN GIẢN (concat/gated) đã đủ; routing phức tạp KHÔNG thêm "
  "giá trị có ý nghĩa — do tín hiệu xung đột yếu/nhiễu và emoji đa nghĩa.")
b("q_j thực nghiệm tiếng Việt: “:)”/“🙂”→anger (mỉa mai), khác hẳn lexicon phương Tây.")
b("Phân tích nhập nhằng amusement–sarcasm + entropy emoji (định lượng tính đa nghĩa).")
b("Audit dữ liệu minh bạch: conflict 22%, nhiễu nhãn mapping 2,3%, bất đối xứng cực ~53%.")
b("Tính giải thích được: aux căn đúng router theo congruence; causal/faithfulness cho thấy mô hình "
  "dùng marker; negative result về β (emoji-dominance không do dữ liệu đòi).")

h2("6.7 Bộ test mỉa mai khách quan: hiệu ứng “emoji hai lưỡi”")
p("Chúng tôi dựng một bộ test mỉa mai khách quan (model-free): mẫu test có emoji rõ cực DƯƠNG "
  "(margin của q_j thực nghiệm ≥ ngưỡng) nhưng nhãn vàng ÂM. Trên bộ này, ngoài accuracy, chúng tôi "
  "đo polarity-accuracy = tỉ lệ model dự đoán đúng cực ÂM (không bị emoji dương đánh lừa) = "
  "“độ kháng mỉa mai”. Đánh giá trên predictions 5-seed đã lưu (không train thêm).")
p("Bảng 5. Polarity-accuracy (kháng mỉa mai) theo độ rõ của emoji (A100 5-seed).", ("", "i"))
table(["margin (độ rõ emoji+)", "n", "B1 text-only", "B2 concat", "B4 cross-attn", "CARE_full"],
      [["≥ 0.0", 104, "0.775", "0.815", "0.706", "0.771"],
       ["≥ 0.1", 75, "0.829", "0.840", "0.715", "0.771"],
       ["≥ 0.3", 36, "0.856", "0.833", "0.717", "0.733"],
       ["≥ 0.5", 16, "0.863", "0.800", "0.650", "0.700"]])
p(("Phát hiện: ", "b"), "trên mỉa mai, biểu tượng cảm xúc là một CÁI BẪY. Model chỉ-text (B1) — "
  "không nhìn emoji — kháng mỉa mai tốt nhất; các model hợp nhất đưa emoji dương vào nên bị kéo về "
  "cực dương và kháng kém hơn (B4 cross-attention, tích hợp nặng nhất, kém nhất). Hiệu ứng ĐƠN ĐIỆU "
  "theo độ rõ của emoji: emoji dương càng rõ (bẫy càng mạnh), B1 càng vượt lên còn B4 càng tụt. "
  "Kết hợp với Mục 6.2 (emoji giúp ở mẫu đồng thuận), điều này lý giải vì sao fusion ≈ text-only ở "
  "overall: ", ("emoji là tín hiệu hai lưỡi — giúp khi đồng thuận, gài bẫy khi mỉa mai", "b"),
  ", và không kiến trúc fusion nào trong khảo sát gỡ được bẫy này.")

h2("6.8 Emoji-như-văn-bản-được-nhân-bản: baseline đơn giản ngang fusion")
p("Một phát hiện kỹ thuật: tokenizer của PhoBERT ", ("đại diện được emoji", "b"),
  " (vd “😂”→[“😂”], “buồn 😢 quá”→[“buồn”,“😢”,“quá”]) — KHÔNG biến thành <unk>. Do đó việc tách "
  "emoji thành một thể thức riêng (qua q_j) là không cần thiết; có thể giữ emoji trong luồng văn bản. "
  "Ta thử một phương án ĐƠN GIẢN thay cho routing: giữ emoji trong text và NHÂN BẢN mỗi emoji K lần "
  "để tăng sức nặng, rồi phân loại text thông thường.")
p("Bảng 6. Emoji-amplified TF-IDF + Logistic Regression theo số lần nhân bản K.", ("", "i"))
table(["K", "test macro-F1", "acc đồng thuận", "acc mỉa mai"],
      [["0 (chỉ text)", "0.410", "0.671", "0.621"],
       ["1", "0.478", "0.786", "0.600"],
       ["3", "0.493", "0.832", "0.524"],
       ["5", "0.498", "0.832", "0.524"],
       ["10", "0.500", "0.834", "0.524"]])
p(("Hai kết luận: ", "b"), "(1) chỉ với TF-IDF + nhân bản emoji đã đạt macro-F1 = 0.500, ",
  ("NGANG CARE-Fusion (0.506)", "b"), " — một phương pháp tầm thường sánh ngang kiến trúc routing "
  "phức tạp, củng cố thông điệp “hợp nhất đơn giản là đủ”. (2) Tham số K phơi bày trực tiếp hiệu ứng "
  "EMOJI HAI LƯỠI: K tăng → đồng thuận tăng mạnh (0.67→0.83) nhưng mỉa mai giảm (0.62→0.52). "
  "Bản PhoBERT (giữ emoji + nhân bản K=3, 1 seed): macro-F1 = 0.480 — nằm trong cùng dải với các "
  "baseline và CARE-Fusion (~0.48–0.51, trong khoảng nhiễu), xác nhận phương án đơn giản cạnh tranh "
  "được với fusion phức tạp (cần đa seed để khẳng định chắc).")

# ============================== 7. LIMITATIONS ==============================
h1("7. Hạn chế")
b("Một dataset / một ngôn ngữ → cần dataset thứ 2 (vd UIT-VSMEC) để khẳng định tổng quát.")
b("Chỉ văn bản có emoji → không đại diện toàn bộ văn bản cảm xúc tiếng Việt (phạm vi có chủ đích).")
b("neutral cực thưa (val 4 / test 6) → báo cáo riêng kèm cảnh báo low-power.")
b("Tầng conflict tự động (proxy δ) chưa kiểm chứng bằng người (~300 mẫu) như đề xuất F2.")
b("Ablation −q (lexicon phương Tây) cần một emoji-sentiment lexicon ngoài để hoàn tất.")
b("Không có trường nền tảng → không đánh giá cross-platform.")

# ============================== 8. CONCLUSION ==============================
h1("8. Kết luận")
p("Biểu tượng cảm xúc cải thiện rõ nhận diện cảm xúc tiếng Việt. Tuy nhiên, trên bộ dữ liệu này, "
  "khó khăn cốt lõi là sự nhập nhằng amusement–sarcasm và tính đa nghĩa của emoji, khiến hợp nhất "
  "đơn giản đã cạnh tranh với các kiến trúc phức tạp hơn. Chúng tôi đóng góp một tập dữ liệu chuyên "
  "biệt, một bộ phân tích định lượng (entropy emoji, q_j tiếng Việt, lát mỉa mai), và một khung "
  "CARE-Fusion có thể giải thích được. Kiểm định 5-seed cho thấy định tuyến theo chế độ không vượt "
  "hợp nhất đơn giản có ý nghĩa trên bộ này — một kết quả trung thực, đồng thời mở ra hướng tương "
  "lai: dataset giàu mỉa mai hơn, gán tay tầng xung đột, và cơ chế irony-aware dùng ngữ cảnh/tri thức.")

doc.add_page_break()
# ============================== APPENDIX: LOG ==============================
h1("Phụ lục A. Nhật ký thí nghiệm & quyết định")
h2("A.1 Hạ tầng")
b("Dự án: C:\\Dev\\care-fusion-sentiment-detection-vi; GitHub dtlong1979 (public).")
b("Local: Python 3.11, torch 2.6+cu124, RTX 3050 6GB (pilot); Java 21 cho VnCoreNLP.")
b("Chạy thật: Colab Pro A100-40GB (~5,3 CU/giờ).")
h2("A.2 Tiến trình & phát hiện chính")
n("Phase 1: tiền xử lý (NFC, trích marker longest-match, word-segment), q_j + đồ thị PMI. Đã verify "
  "trên 6.813 mẫu; 11/11 unit test marker.")
n("Phase 2: model CARE-Fusion, losses (focal-CB + routing + counterfactual), engine multi-seed, "
  "baselines B0–B4, OOF→weak labels.")
n("Phase 3: evaluate F1–F4 + G; F5 (độ nhạy tie-break) qua scripts/f5_sensitivity.py.")
n("Tier-0 (TF-IDF): text 0.41 → text+marker 0.50 → marker có ích.")
n("Pilot 3050: CARE ~0.53 > B1 ~0.47; nhưng CARE ≈ B4 (1 seed).")
n("Chẩn đoán: lỗi p_text không giám sát → router lệch; vá bằng L_aux → router căn đúng δ.")
n("Negative result: β trung tính phẳng → emoji-dominance không do dữ liệu đòi.")
n("Audit dữ liệu: conflict 21,9%; multigroup 31,69%; nhiễu nhãn 2,3%; bất đối xứng cực ~53%.")
n("Phân tích emoji: 18/113 marker nhập nhằng; “:)”/“🙂”→anger (đặc thù tiếng Việt).")
n("Confidence-gating: 1-seed lát sarcasm 0.582 (hứa hẹn) — đang kiểm 5-seed.")
n("A100 5-seed (config cũ): CARE không vượt baseline có ý nghĩa; B2 concat tốt nhất trên lát sarcasm.")
h2("A.3 Quy ước chạy")
b("Profiles: --profile {full, pilot (1 seed), pilot3 (3 seed), pilot5 (5 seed), smoke}.")
b("scripts: data_audit, q_entropy, conflict_asymmetry, list_conflict_cases, f5_sensitivity, "
  "slice_summary, inspect_router, compute_multigroup, quick_tfidf.")

# ============================== REFERENCES ==============================
h1("Tài liệu tham khảo (sơ bộ — cần định dạng chuẩn + xác minh)")
refs = [
    "Novak, P. K., Smailović, J., Sluban, B., Mozetič, I. (2015). Sentiment of Emojis. PLOS ONE. arXiv:1509.07761.",
    "Eisner, B. và cộng sự (2016). emoji2vec: Learning Emoji Representations from their Description. W-NUT.",
    "Felbo, B. và cộng sự (2017). Using millions of emoji occurrences... (DeepMoji). EMNLP. arXiv:1708.00524.",
    "Wang và cộng sự (2024). EMFSA: emoji-based multimodal fusion for sentiment analysis. PLOS ONE 19(10).",
    "Chauhan, D. S. và cộng sự (2022). Gated emoji-aware multimodal sarcasm. Knowledge-Based Systems 257.",
    "Zhu, Peng, Zhang (2025). ECCA: emoji contrastive attention for sarcasm. ICIC, Springer CCIS. [cần xác minh]",
    "Nguyen, D. Q., Nguyen, A. T. (2020). PhoBERT: Pre-trained language models for Vietnamese. EMNLP Findings.",
    "Nguyen và cộng sự (2023). ViSoBERT: pre-trained LM for Vietnamese social media. EMNLP. arXiv:2310.11166.",
    "Thin, D. V. và cộng sự (2023). Vietnamese sentiment/emotion với PLMs (khảo sát). ACM TALLIP.",
    "Hung, T. Q. và cộng sự (2026). ViGoEmotions. EACL. [cần xác minh]",
    "Tran Thi Dung và cộng sự (2024). Nhận diện cảm xúc tiếng Việt kèm emoji. HNUE J. Sci. [cần xác minh]",
    "Cui, Y. và cộng sự (2019). Class-Balanced Loss Based on Effective Number of Samples. CVPR.",
    "Demszky, D. và cộng sự (2020). GoEmotions. ACL.",
]
for ref in refs:
    doc.add_paragraph(ref, style="List Number")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print("Saved:", OUT)
